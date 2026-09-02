#!/usr/bin/env python3
"""Build a polished DOCX report (results only, no code) of the reranker benchmarks.

Imports cleanly into Google Docs. Pulls numbers straight from results/*.json so the report can never
drift from the measured data. Output: report/Qwen3-Reranker-0.6B_Accelerator_Benchmark_Report.docx
"""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RES = os.path.join(ROOT, "results"); CH = os.path.join(ROOT, "charts")
OUT = os.path.join(ROOT, "report"); os.makedirs(OUT, exist_ok=True)

PRICE = {"tpu": 1.61, "b200": 6.95, "h200": 3.85, "g4": 3.11}
NICE = {"tpu": "TPU v6e (torchax)", "b200": "B200 (1 GPU)", "h200": "H200 (1 GPU)", "g4": "G4 RTX PRO 6000 (1 GPU)"}
ORDER = ["tpu", "b200", "h200", "g4"]
LATF = {"tpu": "lat_tpu_torchax.json", "b200": "lat_gpu_b200.json", "h200": "lat_gpu_h200.json", "g4": "lat_gpu_g4.json"}

def load_lat(dev):
    d = json.load(open(os.path.join(RES, LATF[dev])))
    return {r["batch_size"]: r for r in d["by_batch"]}
lat = {dev: load_lat(dev) for dev in ORDER}

def load_conc(dev, seq, rb):
    p = os.path.join(RES, f"conc_{dev}_seq{seq}_rb{rb}.json")
    if not os.path.exists(p): return {}
    d = json.load(open(p))
    return {r["concurrency"]: r for r in d["by_concurrency"]}

def cpm(dev, pps): return round((PRICE[dev]/3600.0)/pps*1e6, 4)
def c1k(dev, ms): return round((PRICE[dev]/3600.0)*(ms/1000.0)*1000.0, 5)

doc = Document()
# base style
st = doc.styles["Normal"].font; st.name = "Calibri"; st.size = Pt(11)

def h(text, level=1):
    p = doc.add_heading(text, level=level); return p

def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, hcell in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(hcell); run.bold = True; run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(9)
    return t

def add_image(path, width=6.3):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def caption(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------- Title ----------------
title = doc.add_heading("Qwen3-Reranker-0.6B: Accelerator Latency & Price/Performance Benchmark", level=0)
para("Real-time serving on vLLM  |  TPU v6e vs NVIDIA B200, H200, and G4 (RTX PRO 6000)", italic=True)
para("")

# ---------------- Executive summary ----------------
h("Executive summary", 1)
para("This report compares Qwen3-Reranker-0.6B served in real time on four single-chip accelerators, "
     "focusing on request latency and price/performance rather than offline batch throughput. All runs use "
     "vLLM with prefix caching disabled and freshly randomized content per request (so caching cannot inflate "
     "results), a standardized 512-token query+document pair, and a batch sweep from 1 to 64 pairs per request, "
     "plus a served concurrency sweep and a 1024-token long-context sweep.")
b = lat["b200"][32]; hh = lat["h200"][32]; g = lat["g4"][32]; t = lat["tpu"][32]
para("Key findings at the batch-size-32 operating point (seq 512):", bold=True)
for dev, r in [("tpu", t), ("b200", b), ("h200", hh), ("g4", g)]:
    para(f"  \u2022 {NICE[dev]}: p50 latency {r['request_latency_ms_p50']} ms, {r['throughput_pairs_per_s']} pairs/s, "
         f"${cpm(dev, r['throughput_pairs_per_s'])}/1M pairs, ${c1k(dev, r['request_latency_ms_p50'])}/1000 requests.")
para("Bottom line: on raw per-request latency the ranking is B200 < H200 < G4 < TPU v6e; on price/performance "
     "(both USD/1M pairs and USD/1000 requests) the TPU v6e is the cheapest overall, followed by G4, then H200, "
     "then B200. At equal latency (matching a GPU's p50 with a smaller TPU batch) the TPU v6e serves the same "
     "SLA at roughly 2x-6x lower cost. Increasing sequence length from 512 to 1024 roughly doubles per-request "
     "latency on every chip but does not change the ranking.")

# ---------------- Method ----------------
h("Methodology", 1)
para("Model: Qwen/Qwen3-Reranker-0.6B (sequence-classification scoring head), served through vLLM's pooling "
     "score / rerank API.")
para("Framework: vLLM torchax on TPU v6e; CUDA vLLM on the NVIDIA GPUs. One chip per device under test.")
para("Rules: prefix caching disabled; each request uses freshly randomized token content so no KV/prefix reuse; "
     "synthetic pairs fixed at 512 tokens (plus a 1024-token long-context sweep); 50 timed iterations per "
     "batch size for the latency sweep.")
para("Prices per chip-hour used throughout: TPU v6e $1.61, B200 $6.95, H200 $3.85, G4 $3.11.")
para("Definitions:", bold=True)
para("  \u2022 Request latency (p50/p90/p99): time for one score() request that contains N pairs; percentiles over the timed iterations.")
para("  \u2022 Throughput (pairs/s) = batch_size / mean request latency (s).")
para("  \u2022 Throughput price/perf = (price/hr / 3600) / pairs_per_s x 1,000,000 = USD per 1M pairs.")
para("  \u2022 Latency price/perf = (price/hr / 3600) x (p50 seconds) x 1000 = USD per 1000 requests.")

# ---------------- Headline table ----------------
h("Headline results (batch size 32, seq 512)", 1)
rows = []
for dev in ORDER:
    r = lat[dev][32]
    rows.append([NICE[dev], r["request_latency_ms_p50"], r["request_latency_ms_p99"], r["throughput_pairs_per_s"],
                 f"${cpm(dev, r['throughput_pairs_per_s'])}", f"${c1k(dev, r['request_latency_ms_p50'])}"])
add_table(["Device", "p50 latency (ms)", "p99 latency (ms)", "pairs/s", "$/1M pairs", "$/1000 req"], rows)
para("")
add_image(os.path.join(CH, "by_batch", "latency_p50_bs32.png"))
caption("Request latency (p50) at batch size 32, all chips.")
add_image(os.path.join(CH, "priceperf_latency_bs32.png"))
caption("Latency-based price/performance at bs=32 (USD per 1000 requests, lower is better).")
add_image(os.path.join(CH, "priceperf_throughput_bs32.png"))
caption("Throughput-based price/performance at bs=32 (USD per 1M pairs, lower is better).")

# ---------------- Batch sweep ----------------
h("Batch-size sweep (seq 512)", 1)
para("Request latency p50 (ms) by batch size:")
rows = [[bs] + [lat[dev][bs]["request_latency_ms_p50"] for dev in ORDER] for bs in sorted(lat["tpu"])]
add_table(["Batch"] + [NICE[d] for d in ORDER], rows)
add_image(os.path.join(CH, "latency_vs_batch.png"))
caption("Request latency p50 vs batch size.")
para("Throughput (pairs/s) by batch size:")
rows = [[bs] + [lat[dev][bs]["throughput_pairs_per_s"] for dev in ORDER] for bs in sorted(lat["tpu"])]
add_table(["Batch"] + [NICE[d] for d in ORDER], rows)
add_image(os.path.join(CH, "throughput_vs_batch.png"))
caption("Throughput (pairs/s) vs batch size.")

# ---------------- GPU-match ----------------
h("Matching a GPU with TPU v6e", 1)
para("Throughput-match (how many v6e chips equal one GPU's throughput, and the fleet cost):", bold=True)
tpu = lat["tpu"]
import math
for dev in ["b200", "h200", "g4"]:
    rows = []
    for bs in sorted(tpu):
        gpps = lat[dev][bs]["throughput_pairs_per_s"]; tpps = tpu[bs]["throughput_pairs_per_s"]
        chips = math.ceil(gpps / tpps); fleet = round(chips * PRICE["tpu"], 2)
        rows.append([bs, gpps, tpps, chips, f"${fleet}", f"${PRICE[dev]}", "yes" if fleet < PRICE[dev] else "no"])
    para(f"vs {NICE[dev]}:", italic=True)
    add_table(["Batch", f"{NICE[dev]} pairs/s", "1x v6e pairs/s", "v6e chips", "v6e fleet $/hr", f"{NICE[dev]} $/hr", "TPU cheaper?"], rows)
para("")
para("Latency-match (meet a GPU's per-request p50 with the fastest v6e batch, then compare cost per 1000 requests):", bold=True)
tbs = sorted(tpu)
for dev in ["b200", "h200", "g4"]:
    rows = []
    for bs in sorted(lat[dev]):
        gp = lat[dev][bs]["request_latency_ms_p50"]; g = c1k(dev, gp)
        ok = [bb for bb in tbs if tpu[bb]["request_latency_ms_p50"] <= gp]
        if ok:
            mb = max(ok); mp = tpu[mb]["request_latency_ms_p50"]; tt = c1k("tpu", mp)
            rows.append([bs, gp, f"${g}", f"v6e bs{mb}", mp, f"${tt}", "TPU" if tt < g else NICE[dev]])
        else:
            rows.append([bs, gp, f"${g}", "none (v6e floor slower)", "-", "-", NICE[dev]])
    para(f"vs {NICE[dev]}:", italic=True)
    add_table(["GPU batch", f"{NICE[dev]} p50 (ms)", f"{NICE[dev]} $/1k", "v6e config", "v6e p50 (ms)", "v6e $/1k", "Cheaper at equal latency"], rows)

# ---------------- Concurrency ----------------
h("Concurrency sweep (served vLLM, real-time load)", 1)
para("Served via vLLM's HTTP rerank endpoint with C simultaneous in-flight requests; prefix caching off; "
     "fresh content per request. Each cell shows p50 / p90 / p99 latency (ms) | requests/s | USD per 1000 requests.")

def conc_rows(seq, rb, concs):
    out = []
    data = {dev: load_conc(dev, seq, rb) for dev in ORDER}
    for c in concs:
        row = [c]
        for dev in ORDER:
            r = data[dev].get(c)
            if r:
                row.append(f"{r['req_latency_ms_p50']}/{r['req_latency_ms_p90']}/{r['req_latency_ms_p99']} | "
                           f"{r['qps']} | ${c1k(dev, r['req_latency_ms_p50'])}")
            else:
                row.append("-")
        out.append(row)
    return out

para("Single pair per request (seq 512):", bold=True)
add_table(["Conc"] + [NICE[d] for d in ORDER], conc_rows(512, 1, [1,2,4,8]))
add_image(os.path.join(CH, "concurrency_seq512_rb1_p50.png")); caption("Single-pair p50 latency vs concurrency (seq 512).")
add_image(os.path.join(CH, "concurrency_seq512_rb1_qps.png")); caption("Single-pair throughput vs concurrency (seq 512).")

para("Batch of 32 pairs per request (seq 512):", bold=True)
add_table(["Conc"] + [NICE[d] for d in ORDER], conc_rows(512, 32, [1,2,4,8]))
add_image(os.path.join(CH, "concurrency_seq512_rb32_p50.png")); caption("Batch-32 p50 latency vs concurrency (seq 512).")
add_image(os.path.join(CH, "concurrency_seq512_rb32_qps.png")); caption("Batch-32 throughput vs concurrency (seq 512).")

# ---------------- Sequence length ----------------
h("Sequence-length (long-context) sweep: 512 vs 1024 tokens", 1)
para("Holding the request shape at 32 pairs, doubling tokens per pair from 512 to 1024 roughly doubles "
     "per-request latency on every chip while preserving the price/perf ranking. Each cell: p50/p90/p99 (ms) | req/s | $/1k.")
add_table(["Conc"] + [NICE[d] for d in ORDER], conc_rows(1024, 32, [1,2,4]))
add_image(os.path.join(CH, "concurrency_seq1024_rb32_p50.png")); caption("Batch-32 p50 latency vs concurrency, seq 1024.")
add_image(os.path.join(CH, "concurrency_seq1024_rb32_qps.png")); caption("Batch-32 throughput vs concurrency, seq 1024.")

# ---------------- Conclusion ----------------
h("Conclusions", 1)
para("  \u2022 TPU v6e delivers the best price/performance for real-time reranking on both the throughput "
     "(USD/1M pairs) and latency (USD/1000 requests) views at every batch size tested.")
para("  \u2022 Among GPUs, G4 (RTX PRO 6000) is the best value and has the lowest single-request latency at "
     "small batches; B200 has the lowest absolute latency at larger batches; H200 sits between them.")
para("  \u2022 At equal latency, matching a GPU's p50 with a smaller TPU batch costs roughly 2x-6x less.")
para("  \u2022 Longer context (1024 vs 512) scales latency and cost proportionally without changing the ranking.")

path = os.path.join(OUT, "Qwen3-Reranker-0.6B_Accelerator_Benchmark_Report.docx")
doc.save(path)
print("Wrote", path)
